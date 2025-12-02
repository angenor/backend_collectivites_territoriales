"""
Export service for generating Excel and Word documents.
Generates administrative account reports in various formats.
"""

from datetime import datetime
from decimal import Decimal
from io import BytesIO
from typing import Optional

from openpyxl import Workbook
from openpyxl.styles import (
    Alignment,
    Border,
    Font,
    PatternFill,
    Side,
)
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.schemas.tableau import (
    TableauComplet,
    TableauRecettes,
    TableauDepenses,
    SectionTableauRecettes,
    SectionTableauDepenses,
)


def format_montant(montant: Optional[Decimal], with_symbol: bool = False) -> str:
    """Format amount in Ariary (MGA) with thousands separator."""
    if montant is None:
        return ""
    value = float(montant)
    formatted = f"{value:,.0f}".replace(",", " ")
    if with_symbol:
        return f"{formatted} MGA"
    return formatted


def format_taux(taux: Optional[Decimal]) -> str:
    """Format percentage."""
    if taux is None:
        return ""
    return f"{float(taux):.1f}%"


class ExcelExportService:
    """Service for generating Excel exports."""

    def __init__(self):
        # Styles
        self.header_font = Font(bold=True, size=11)
        self.title_font = Font(bold=True, size=14)
        self.section_font = Font(bold=True, size=12, color="FFFFFF")
        self.normal_font = Font(size=10)

        self.header_fill = PatternFill(
            start_color="4472C4", end_color="4472C4", fill_type="solid"
        )
        self.section_fill = PatternFill(
            start_color="5B9BD5", end_color="5B9BD5", fill_type="solid"
        )
        self.total_fill = PatternFill(
            start_color="D9E2F3", end_color="D9E2F3", fill_type="solid"
        )
        self.level2_fill = PatternFill(
            start_color="F2F2F2", end_color="F2F2F2", fill_type="solid"
        )

        self.thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        self.center_align = Alignment(horizontal="center", vertical="center")
        self.right_align = Alignment(horizontal="right", vertical="center")
        self.left_align = Alignment(horizontal="left", vertical="center")

    def _set_cell_style(
        self,
        cell,
        font=None,
        fill=None,
        alignment=None,
        border=True,
    ):
        """Apply styles to a cell."""
        if font:
            cell.font = font
        if fill:
            cell.fill = fill
        if alignment:
            cell.alignment = alignment
        if border:
            cell.border = self.thin_border

    def _add_header_row(self, ws, row: int, columns: list, col_start: int = 1):
        """Add a header row with styling."""
        for i, col_name in enumerate(columns):
            cell = ws.cell(row=row, column=col_start + i, value=col_name)
            self._set_cell_style(
                cell,
                font=self.header_font,
                fill=self.header_fill,
                alignment=self.center_align,
            )
            cell.font = Font(bold=True, color="FFFFFF")

    def _auto_width(self, ws, min_width: int = 10, max_width: int = 50):
        """Auto-adjust column widths."""
        for column_cells in ws.columns:
            length = max(
                min_width,
                min(
                    max_width,
                    max(len(str(cell.value or "")) for cell in column_cells) + 2,
                ),
            )
            ws.column_dimensions[
                get_column_letter(column_cells[0].column)
            ].width = length

    def generate_tableau_complet(self, tableau: TableauComplet) -> BytesIO:
        """Generate complete administrative account Excel file."""
        wb = Workbook()

        # Sheet 1: Recettes
        ws_recettes = wb.active
        ws_recettes.title = "Recettes"
        self._write_recettes_sheet(ws_recettes, tableau.recettes, tableau)

        # Sheet 2: Depenses
        ws_depenses = wb.create_sheet("Dépenses")
        self._write_depenses_sheet(ws_depenses, tableau.depenses, tableau)

        # Sheet 3: Equilibre
        ws_equilibre = wb.create_sheet("Équilibre")
        self._write_equilibre_sheet(ws_equilibre, tableau)

        # Save to BytesIO
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        return output

    def generate_recettes_only(self, recettes: TableauRecettes, commune_info: dict) -> BytesIO:
        """Generate Excel file with only receipts table."""
        wb = Workbook()
        ws = wb.active
        ws.title = "Recettes"

        # Create a minimal tableau info
        class MinimalInfo:
            commune_nom = commune_info.get("commune_nom", "")
            commune_code = commune_info.get("commune_code", "")
            region_nom = commune_info.get("region_nom", "")
            province_nom = commune_info.get("province_nom", "")
            exercice_annee = recettes.exercice_annee

        self._write_recettes_sheet(ws, recettes, MinimalInfo())

        output = BytesIO()
        wb.save(output)
        output.seek(0)
        return output

    def generate_depenses_only(self, depenses: TableauDepenses, commune_info: dict) -> BytesIO:
        """Generate Excel file with only expenses table."""
        wb = Workbook()
        ws = wb.active
        ws.title = "Dépenses"

        class MinimalInfo:
            commune_nom = commune_info.get("commune_nom", "")
            commune_code = commune_info.get("commune_code", "")
            region_nom = commune_info.get("region_nom", "")
            province_nom = commune_info.get("province_nom", "")
            exercice_annee = depenses.exercice_annee

        self._write_depenses_sheet(ws, depenses, MinimalInfo())

        output = BytesIO()
        wb.save(output)
        output.seek(0)
        return output

    def _write_header_info(self, ws, tableau_info, row_start: int = 1) -> int:
        """Write header information and return next available row."""
        ws.cell(row=row_start, column=1, value="COMPTE ADMINISTRATIF")
        ws.cell(row=row_start, column=1).font = self.title_font
        ws.merge_cells(start_row=row_start, start_column=1, end_row=row_start, end_column=4)

        ws.cell(row=row_start + 1, column=1, value=f"Commune: {tableau_info.commune_nom}")
        ws.cell(row=row_start + 1, column=3, value=f"Code: {tableau_info.commune_code}")
        ws.cell(row=row_start + 2, column=1, value=f"Région: {tableau_info.region_nom}")
        ws.cell(row=row_start + 2, column=3, value=f"Province: {tableau_info.province_nom}")
        ws.cell(row=row_start + 3, column=1, value=f"Exercice: {tableau_info.exercice_annee}")
        ws.cell(
            row=row_start + 3,
            column=3,
            value=f"Date: {datetime.now().strftime('%d/%m/%Y')}",
        )

        return row_start + 5

    def _write_recettes_sheet(self, ws, recettes: TableauRecettes, tableau_info) -> None:
        """Write receipts data to worksheet."""
        row = self._write_header_info(ws, tableau_info)

        # Title
        ws.cell(row=row, column=1, value="TABLEAU DES RECETTES")
        ws.cell(row=row, column=1).font = self.title_font
        row += 2

        # Headers
        headers = [
            "Code",
            "Intitulé",
            "Budget Primitif",
            "Budget Additionnel",
            "Modifications",
            "Prévisions Définitives",
            "OR Admis",
            "Recouvrement",
            "Reste à Recouvrer",
            "Taux Exécution",
        ]
        self._add_header_row(ws, row, headers)
        row += 1

        # Data for each section
        for section in recettes.sections:
            # Section header
            ws.cell(row=row, column=1, value=section.titre)
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=10)
            self._set_cell_style(
                ws.cell(row=row, column=1),
                font=self.section_font,
                fill=self.section_fill,
                alignment=self.left_align,
            )
            row += 1

            # Data rows
            for ligne in section.lignes:
                ws.cell(row=row, column=1, value=ligne.code)
                ws.cell(row=row, column=2, value=ligne.intitule)
                ws.cell(row=row, column=3, value=format_montant(ligne.budget_primitif))
                ws.cell(row=row, column=4, value=format_montant(ligne.budget_additionnel))
                ws.cell(row=row, column=5, value=format_montant(ligne.modifications))
                ws.cell(row=row, column=6, value=format_montant(ligne.previsions_definitives))
                ws.cell(row=row, column=7, value=format_montant(ligne.or_admis))
                ws.cell(row=row, column=8, value=format_montant(ligne.recouvrement))
                ws.cell(row=row, column=9, value=format_montant(ligne.reste_a_recouvrer))
                ws.cell(row=row, column=10, value=format_taux(ligne.taux_execution))

                # Style based on level
                fill = self.level2_fill if ligne.niveau > 1 else None
                for col in range(1, 11):
                    cell = ws.cell(row=row, column=col)
                    align = self.right_align if col > 2 else self.left_align
                    self._set_cell_style(cell, fill=fill, alignment=align)
                    if ligne.niveau == 1:
                        cell.font = Font(bold=True, size=10)

                row += 1

            # Section total
            ws.cell(row=row, column=1, value="")
            ws.cell(row=row, column=2, value=f"TOTAL {section.titre}")
            ws.cell(row=row, column=3, value=format_montant(section.total_budget_primitif))
            ws.cell(row=row, column=4, value=format_montant(section.total_budget_additionnel))
            ws.cell(row=row, column=5, value=format_montant(section.total_modifications))
            ws.cell(row=row, column=6, value=format_montant(section.total_previsions_definitives))
            ws.cell(row=row, column=7, value=format_montant(section.total_or_admis))
            ws.cell(row=row, column=8, value=format_montant(section.total_recouvrement))
            ws.cell(row=row, column=9, value=format_montant(section.total_reste_a_recouvrer))
            ws.cell(row=row, column=10, value=format_taux(section.taux_execution_global))

            for col in range(1, 11):
                self._set_cell_style(
                    ws.cell(row=row, column=col),
                    font=Font(bold=True),
                    fill=self.total_fill,
                    alignment=self.right_align if col > 2 else self.left_align,
                )
            row += 2

        # Grand total
        ws.cell(row=row, column=2, value="TOTAL GÉNÉRAL")
        ws.cell(row=row, column=6, value=format_montant(recettes.total_general_previsions))
        ws.cell(row=row, column=7, value=format_montant(recettes.total_general_or_admis))
        ws.cell(row=row, column=8, value=format_montant(recettes.total_general_recouvrement))
        ws.cell(row=row, column=10, value=format_taux(recettes.taux_execution_global))

        for col in range(1, 11):
            self._set_cell_style(
                ws.cell(row=row, column=col),
                font=Font(bold=True, size=11),
                fill=self.header_fill,
                alignment=self.right_align if col > 2 else self.left_align,
            )
            ws.cell(row=row, column=col).font = Font(bold=True, color="FFFFFF")

        self._auto_width(ws)

    def _write_depenses_sheet(self, ws, depenses: TableauDepenses, tableau_info) -> None:
        """Write expenses data to worksheet."""
        row = self._write_header_info(ws, tableau_info)

        # Title
        ws.cell(row=row, column=1, value="TABLEAU DES DÉPENSES")
        ws.cell(row=row, column=1).font = self.title_font
        row += 2

        # Headers
        headers = [
            "Code",
            "Intitulé",
            "Budget Primitif",
            "Budget Additionnel",
            "Modifications",
            "Prévisions Définitives",
            "Engagement",
            "Mandat Admis",
            "Paiement",
            "Reste à Payer",
            "Taux Exécution",
        ]
        self._add_header_row(ws, row, headers)
        row += 1

        # Data for each section
        for section in depenses.sections:
            # Section header
            ws.cell(row=row, column=1, value=section.titre)
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=11)
            self._set_cell_style(
                ws.cell(row=row, column=1),
                font=self.section_font,
                fill=self.section_fill,
                alignment=self.left_align,
            )
            row += 1

            # Data rows
            for ligne in section.lignes:
                ws.cell(row=row, column=1, value=ligne.code)
                ws.cell(row=row, column=2, value=ligne.intitule)
                ws.cell(row=row, column=3, value=format_montant(ligne.budget_primitif))
                ws.cell(row=row, column=4, value=format_montant(ligne.budget_additionnel))
                ws.cell(row=row, column=5, value=format_montant(ligne.modifications))
                ws.cell(row=row, column=6, value=format_montant(ligne.previsions_definitives))
                ws.cell(row=row, column=7, value=format_montant(ligne.engagement))
                ws.cell(row=row, column=8, value=format_montant(ligne.mandat_admis))
                ws.cell(row=row, column=9, value=format_montant(ligne.paiement))
                ws.cell(row=row, column=10, value=format_montant(ligne.reste_a_payer))
                ws.cell(row=row, column=11, value=format_taux(ligne.taux_execution))

                # Style based on level
                fill = self.level2_fill if ligne.niveau > 1 else None
                for col in range(1, 12):
                    cell = ws.cell(row=row, column=col)
                    align = self.right_align if col > 2 else self.left_align
                    self._set_cell_style(cell, fill=fill, alignment=align)
                    if ligne.niveau == 1:
                        cell.font = Font(bold=True, size=10)

                row += 1

            # Section total
            ws.cell(row=row, column=1, value="")
            ws.cell(row=row, column=2, value=f"TOTAL {section.titre}")
            ws.cell(row=row, column=3, value=format_montant(section.total_budget_primitif))
            ws.cell(row=row, column=4, value=format_montant(section.total_budget_additionnel))
            ws.cell(row=row, column=5, value=format_montant(section.total_modifications))
            ws.cell(row=row, column=6, value=format_montant(section.total_previsions_definitives))
            ws.cell(row=row, column=7, value=format_montant(section.total_engagement))
            ws.cell(row=row, column=8, value=format_montant(section.total_mandat_admis))
            ws.cell(row=row, column=9, value=format_montant(section.total_paiement))
            ws.cell(row=row, column=10, value=format_montant(section.total_reste_a_payer))
            ws.cell(row=row, column=11, value=format_taux(section.taux_execution_global))

            for col in range(1, 12):
                self._set_cell_style(
                    ws.cell(row=row, column=col),
                    font=Font(bold=True),
                    fill=self.total_fill,
                    alignment=self.right_align if col > 2 else self.left_align,
                )
            row += 2

        # Grand total
        ws.cell(row=row, column=2, value="TOTAL GÉNÉRAL")
        ws.cell(row=row, column=6, value=format_montant(depenses.total_general_previsions))
        ws.cell(row=row, column=8, value=format_montant(depenses.total_general_mandat_admis))
        ws.cell(row=row, column=9, value=format_montant(depenses.total_general_paiement))
        ws.cell(row=row, column=11, value=format_taux(depenses.taux_execution_global))

        for col in range(1, 12):
            self._set_cell_style(
                ws.cell(row=row, column=col),
                font=Font(bold=True, size=11),
                fill=self.header_fill,
                alignment=self.right_align if col > 2 else self.left_align,
            )
            ws.cell(row=row, column=col).font = Font(bold=True, color="FFFFFF")

        self._auto_width(ws)

    def _write_equilibre_sheet(self, ws, tableau: TableauComplet) -> None:
        """Write balance table to worksheet."""
        row = self._write_header_info(ws, tableau)

        # Title
        ws.cell(row=row, column=1, value="TABLEAU D'ÉQUILIBRE BUDGÉTAIRE")
        ws.cell(row=row, column=1).font = self.title_font
        row += 2

        # Headers
        headers = [
            "Section",
            "Recettes Prévisions",
            "Recettes Réalisations",
            "Dépenses Prévisions",
            "Dépenses Réalisations",
            "Solde Prévisions",
            "Solde Réalisations",
        ]
        self._add_header_row(ws, row, headers)
        row += 1

        eq = tableau.equilibre

        # Fonctionnement row
        ws.cell(row=row, column=1, value="Section de fonctionnement")
        ws.cell(row=row, column=2, value=format_montant(eq.fonctionnement_recettes_prev))
        ws.cell(row=row, column=3, value=format_montant(eq.fonctionnement_recettes_real))
        ws.cell(row=row, column=4, value=format_montant(eq.fonctionnement_depenses_prev))
        ws.cell(row=row, column=5, value=format_montant(eq.fonctionnement_depenses_real))
        ws.cell(row=row, column=6, value=format_montant(eq.fonctionnement_solde_prev))
        ws.cell(row=row, column=7, value=format_montant(eq.fonctionnement_solde_real))
        for col in range(1, 8):
            self._set_cell_style(
                ws.cell(row=row, column=col),
                alignment=self.right_align if col > 1 else self.left_align,
            )
        row += 1

        # Investissement row
        ws.cell(row=row, column=1, value="Section d'investissement")
        ws.cell(row=row, column=2, value=format_montant(eq.investissement_recettes_prev))
        ws.cell(row=row, column=3, value=format_montant(eq.investissement_recettes_real))
        ws.cell(row=row, column=4, value=format_montant(eq.investissement_depenses_prev))
        ws.cell(row=row, column=5, value=format_montant(eq.investissement_depenses_real))
        ws.cell(row=row, column=6, value=format_montant(eq.investissement_solde_prev))
        ws.cell(row=row, column=7, value=format_montant(eq.investissement_solde_real))
        for col in range(1, 8):
            self._set_cell_style(
                ws.cell(row=row, column=col),
                alignment=self.right_align if col > 1 else self.left_align,
            )
        row += 1

        # Total row
        ws.cell(row=row, column=1, value="TOTAL GÉNÉRAL")
        ws.cell(row=row, column=2, value=format_montant(eq.total_recettes_prev))
        ws.cell(row=row, column=3, value=format_montant(eq.total_recettes_real))
        ws.cell(row=row, column=4, value=format_montant(eq.total_depenses_prev))
        ws.cell(row=row, column=5, value=format_montant(eq.total_depenses_real))
        ws.cell(row=row, column=6, value=format_montant(eq.total_solde_prev))
        ws.cell(row=row, column=7, value=format_montant(eq.total_solde_real))
        for col in range(1, 8):
            self._set_cell_style(
                ws.cell(row=row, column=col),
                font=Font(bold=True),
                fill=self.total_fill,
                alignment=self.right_align if col > 1 else self.left_align,
            )

        self._auto_width(ws)


class WordExportService:
    """Service for generating Word exports."""

    def generate_tableau_complet(self, tableau: TableauComplet) -> BytesIO:
        """Generate complete administrative account Word document."""
        doc = Document()

        # Title
        title = doc.add_heading("COMPTE ADMINISTRATIF", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header info
        doc.add_paragraph(f"Commune: {tableau.commune_nom} ({tableau.commune_code})")
        doc.add_paragraph(f"Région: {tableau.region_nom} - Province: {tableau.province_nom}")
        doc.add_paragraph(f"Exercice: {tableau.exercice_annee}")
        doc.add_paragraph(f"Date de génération: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()

        # Recettes section
        self._add_recettes_table(doc, tableau.recettes)

        doc.add_page_break()

        # Depenses section
        self._add_depenses_table(doc, tableau.depenses)

        doc.add_page_break()

        # Equilibre section
        self._add_equilibre_table(doc, tableau.equilibre)

        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    def _add_recettes_table(self, doc: Document, recettes: TableauRecettes) -> None:
        """Add receipts table to document."""
        doc.add_heading("TABLEAU DES RECETTES", level=1)

        for section in recettes.sections:
            doc.add_heading(section.titre, level=2)

            # Create table
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            header_cells = table.rows[0].cells
            headers = ["Code", "Intitulé", "Prév. Déf.", "OR Admis", "Taux"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True

            # Data rows (only level 1 for Word)
            for ligne in section.lignes:
                if ligne.niveau == 1:
                    row_cells = table.add_row().cells
                    row_cells[0].text = ligne.code
                    row_cells[1].text = ligne.intitule
                    row_cells[2].text = format_montant(ligne.previsions_definitives)
                    row_cells[3].text = format_montant(ligne.or_admis)
                    row_cells[4].text = format_taux(ligne.taux_execution)

            # Total row
            total_row = table.add_row().cells
            total_row[0].text = ""
            total_row[1].text = f"TOTAL {section.titre}"
            total_row[2].text = format_montant(section.total_previsions_definitives)
            total_row[3].text = format_montant(section.total_or_admis)
            total_row[4].text = format_taux(section.taux_execution_global)
            for cell in total_row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            doc.add_paragraph()

        # Grand total
        p = doc.add_paragraph()
        p.add_run("TOTAL GÉNÉRAL RECETTES: ").bold = True
        p.add_run(format_montant(recettes.total_general_or_admis, with_symbol=True))
        p.add_run(f" (Taux: {format_taux(recettes.taux_execution_global)})")

    def _add_depenses_table(self, doc: Document, depenses: TableauDepenses) -> None:
        """Add expenses table to document."""
        doc.add_heading("TABLEAU DES DÉPENSES", level=1)

        for section in depenses.sections:
            doc.add_heading(section.titre, level=2)

            # Create table
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            header_cells = table.rows[0].cells
            headers = ["Code", "Intitulé", "Prév. Déf.", "Mandat Admis", "Taux"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True

            # Data rows (only level 1 for Word)
            for ligne in section.lignes:
                if ligne.niveau == 1:
                    row_cells = table.add_row().cells
                    row_cells[0].text = ligne.code
                    row_cells[1].text = ligne.intitule
                    row_cells[2].text = format_montant(ligne.previsions_definitives)
                    row_cells[3].text = format_montant(ligne.mandat_admis)
                    row_cells[4].text = format_taux(ligne.taux_execution)

            # Total row
            total_row = table.add_row().cells
            total_row[0].text = ""
            total_row[1].text = f"TOTAL {section.titre}"
            total_row[2].text = format_montant(section.total_previsions_definitives)
            total_row[3].text = format_montant(section.total_mandat_admis)
            total_row[4].text = format_taux(section.taux_execution_global)
            for cell in total_row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            doc.add_paragraph()

        # Grand total
        p = doc.add_paragraph()
        p.add_run("TOTAL GÉNÉRAL DÉPENSES: ").bold = True
        p.add_run(format_montant(depenses.total_general_mandat_admis, with_symbol=True))
        p.add_run(f" (Taux: {format_taux(depenses.taux_execution_global)})")

    def _add_equilibre_table(self, doc: Document, equilibre) -> None:
        """Add balance table to document."""
        doc.add_heading("TABLEAU D'ÉQUILIBRE BUDGÉTAIRE", level=1)

        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[0].cells
        headers = ["Section", "Recettes", "Dépenses", "Solde"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Fonctionnement
        row1 = table.add_row().cells
        row1[0].text = "Fonctionnement"
        row1[1].text = format_montant(equilibre.fonctionnement_recettes_real)
        row1[2].text = format_montant(equilibre.fonctionnement_depenses_real)
        row1[3].text = format_montant(equilibre.fonctionnement_solde_real)

        # Investissement
        row2 = table.add_row().cells
        row2[0].text = "Investissement"
        row2[1].text = format_montant(equilibre.investissement_recettes_real)
        row2[2].text = format_montant(equilibre.investissement_depenses_real)
        row2[3].text = format_montant(equilibre.investissement_solde_real)

        # Total
        total_row = table.add_row().cells
        total_row[0].text = "TOTAL"
        total_row[1].text = format_montant(equilibre.total_recettes_real)
        total_row[2].text = format_montant(equilibre.total_depenses_real)
        total_row[3].text = format_montant(equilibre.total_solde_real)
        for cell in total_row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph()

        # Summary
        solde = equilibre.total_solde_real
        status = "excédentaire" if solde >= 0 else "déficitaire"
        p = doc.add_paragraph()
        p.add_run(f"Le budget est {status} avec un solde de: ").bold = True
        p.add_run(format_montant(abs(solde), with_symbol=True))


# Singleton instances
excel_export_service = ExcelExportService()
word_export_service = WordExportService()
